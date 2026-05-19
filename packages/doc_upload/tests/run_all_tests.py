"""
run_all_tests.py — Master test runner for pdf-autofillr-doc-upload
====================================================================
Run this from the doc_upload/ directory:

    python run_all_tests.py

What it does:
  1. Creates all sample data files in data/input/ (xlsx, docx — txt/json/csv/md already bundled)
  2. Copies sample configs to configs/ if not already present
  3. Runs unit tests
  4. Runs functional tests (mocked LLM — no API key needed)
  5. Runs real LLM test (reads DOC_UPLOAD_LLM_API_KEY from .env)
  6. Runs mapper end-to-end test (reads DOC_UPLOAD_PDF_PATH + API key from .env)
  7. Prints a final summary table
"""

import os
import sys
import subprocess
import json
from pathlib import Path
from dotenv import load_dotenv

# ── Load .env ─────────────────────────────────────────────────────────────────
load_dotenv()

ROOT       = Path(__file__).parent
DATA_INPUT = ROOT / "data" / "input"
CONFIGS    = ROOT / "configs"
SCHEMA     = CONFIGS / "form_keys.json"

PYTHON     = sys.executable


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def header(title: str):
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")


def step(msg: str):
    print(f"\n── {msg}")


def ok(msg: str):
    print(f"  ✅  {msg}")


def warn(msg: str):
    print(f"  ⚠️   {msg}")


def info(msg: str):
    print(f"  ℹ️   {msg}")


def run_pytest(label: str, args: list, env: dict = None) -> dict:
    """Run pytest and return {label, passed, failed, skipped, errors, returncode}."""
    cmd = [PYTHON, "-m", "pytest", "--tb=short", "-q", "--no-header"] + args
    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        cwd=str(ROOT),
        env={**os.environ, **(env or {})},
    )
    output = result.stdout + result.stderr

    # Parse counts from pytest summary line
    passed = failed = skipped = errors = 0
    for line in output.splitlines():
        if "passed" in line or "failed" in line or "error" in line or "skipped" in line:
            import re
            p = re.search(r"(\d+) passed",  line); passed  = int(p.group(1)) if p else passed
            f = re.search(r"(\d+) failed",  line); failed  = int(f.group(1)) if f else failed
            s = re.search(r"(\d+) skipped", line); skipped = int(s.group(1)) if s else skipped
            e = re.search(r"(\d+) error",   line); errors  = int(e.group(1)) if e else errors

    # Print output (filter noise)
    for line in output.splitlines():
        skip = any(x in line for x in [
            "PytestConfigWarning", "asyncio_mode", "warn_or_fail",
            "PydanticDeprecated", "ConfigDict", "fields' has been removed",
            "PydanticSerializationUnexpected", "cachedir", "rootdir",
            "configfile", "platform win", "plugins:",
        ])
        if not skip and line.strip():
            print(f"    {line}")

    return {
        "label":   label,
        "passed":  passed,
        "failed":  failed,
        "skipped": skipped,
        "errors":  errors,
        "ok":      failed == 0 and errors == 0,
    }


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — PREFLIGHT CHECKS
# ══════════════════════════════════════════════════════════════════════════════

header("Step 1 — Preflight checks")

# Check configs
if not SCHEMA.exists():
    step("Copying sample configs …")
    result = subprocess.run(
        [PYTHON, "-c",
         "import pdf_autofillr_doc_upload; pdf_autofillr_doc_upload.copy_sample_configs('.')"],
        capture_output=True, text=True, cwd=str(ROOT)
    )
    if result.returncode == 0:
        ok("configs/ created")
    else:
        warn(f"copy_sample_configs failed: {result.stderr.strip()}")
else:
    ok(f"configs/form_keys.json present")

# Check data/input
DATA_INPUT.mkdir(parents=True, exist_ok=True)

# Check env vars
api_key = os.getenv("DOC_UPLOAD_LLM_API_KEY") or os.getenv("OPENAI_API_KEY", "")
pdf_path = os.getenv("DOC_UPLOAD_PDF_PATH", "")
pdf_filler = os.getenv("DOC_UPLOAD_PDF_FILLER", "none")

if api_key and not api_key.startswith("sk-..."):
    ok(f"LLM API key found ({api_key[:12]}...)")
else:
    warn("DOC_UPLOAD_LLM_API_KEY not set in .env — real LLM and mapper e2e tests will be skipped")

if pdf_path and Path(pdf_path).exists():
    ok(f"Blank PDF found: {pdf_path}")
elif pdf_filler == "mapper":
    warn(f"DOC_UPLOAD_PDF_PATH not found: {pdf_path!r} — mapper e2e test will be skipped")
else:
    info("DOC_UPLOAD_PDF_FILLER=none — mapper e2e test will be skipped")

# Check Java (needed for mapper)
java_check = subprocess.run(["java", "-version"], capture_output=True, text=True)
if java_check.returncode == 0:
    java_ver = (java_check.stderr or java_check.stdout).splitlines()[0] if (java_check.stderr or java_check.stdout) else "found"
    ok(f"Java: {java_ver}")
else:
    warn("Java not found on PATH — mapper fill will fail if tested")

# Check C:/Temp (Windows mapper requirement)
if sys.platform == "win32":
    if Path("C:/Temp").exists():
        ok("C:/Temp exists")
    else:
        warn("C:/Temp does not exist — create it: New-Item -ItemType Directory -Path C:\\Temp -Force")


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — CREATE SAMPLE DATA FILES
# ══════════════════════════════════════════════════════════════════════════════

header("Step 2 — Create sample data files")

# ── xlsx ──────────────────────────────────────────────────────────────────────
xlsx_path = DATA_INPUT / "sample_investor.xlsx"
if not xlsx_path.exists():
    step("Creating sample_investor.xlsx …")
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Field", "Value"])
        ws.append(["Investor Full Legal Name", "Robert C. Johnson"])
        ws.append(["Date of Birth", "10/22/1975"])
        ws.append(["Social Security Number", "456-78-9012"])
        ws.append(["Email Address", "robert.johnson@example.com"])
        ws.append(["Telephone Number", "+44-2079460958"])
        ws.append(["Occupation", "Fund Manager"])
        ws.append(["Registered Address Line 1", "88 Wood Street, Floor 3"])
        ws.append(["Registered City", "London"])
        ws.append(["Registered Country", "United Kingdom"])
        ws.append(["Registered Zip", "EC2V 7RS"])
        ws.append(["Bank Name", "HSBC UK Bank plc"])
        ws.append(["Account Name", "RCJ Capital Management"])
        ws.append(["Account Number", "40229744"])
        ws.append(["SWIFT Code", "MIDLGB22"])
        ws.append(["Authorized Signatory", "Robert C. Johnson"])
        ws.append(["Commitment Amount", "1000000"])
        ws.append(["Currency", "GBP"])
        ws.append(["Accredited Investor", "Yes"])
        ws.append(["US Person", "No"])
        ws.append(["Politically Exposed Person", "No"])
        ws.append(["Share Class", "Class B shares"])
        wb.save(str(xlsx_path))
        ok("data/input/sample_investor.xlsx created")
    except ImportError:
        warn("openpyxl not installed — skipping xlsx creation (pip install openpyxl)")
else:
    ok("data/input/sample_investor.xlsx already exists")

# ── docx ──────────────────────────────────────────────────────────────────────
docx_path = DATA_INPUT / "sample_investor.docx"
if not docx_path.exists():
    step("Creating sample_investor.docx …")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Investor Subscription Form", 0)
        doc.add_heading("Personal Information", level=1)
        doc.add_paragraph("Investor Full Legal Name: Test Investor DOCX")
        doc.add_paragraph("Date of Birth: 01/01/1985")
        doc.add_paragraph("Social Security Number: 111-22-3333")
        doc.add_paragraph("Email Address: test.docx@example.com")
        doc.add_paragraph("Telephone Number: +1-3105551234")
        doc.add_paragraph("Occupation: Investment Manager")
        doc.add_heading("Registered Address", level=1)
        doc.add_paragraph("Line 1: 100 Wilshire Blvd, Suite 500")
        doc.add_paragraph("City: Santa Monica")
        doc.add_paragraph("State: CA")
        doc.add_paragraph("Country: USA")
        doc.add_paragraph("Zip: 90401")
        doc.add_heading("Banking / Wiring", level=1)
        doc.add_paragraph("Bank Name: Bank of America NA")
        doc.add_paragraph("Account Name: Test Investor DOCX")
        doc.add_paragraph("Account Number: 9988776655")
        doc.add_paragraph("ABA Routing: 026009593")
        doc.add_paragraph("SWIFT Code: BOFAUS3N")
        doc.add_paragraph("FATF Compliant: Yes")
        doc.add_heading("Investment Details", level=1)
        doc.add_paragraph("Authorized Signatory: Test Investor DOCX")
        doc.add_paragraph("Commitment Amount: $100,000")
        doc.add_paragraph("Currency: USD")
        doc.add_paragraph("Accredited Investor: Yes")
        doc.add_paragraph("US Person: Yes")
        doc.add_paragraph("PEP: No")
        doc.add_paragraph("Share Class: Class A shares")
        doc.save(str(docx_path))
        ok("data/input/sample_investor.docx created")
    except ImportError:
        warn("python-docx not installed — skipping docx creation (pip install python-docx)")
else:
    ok("data/input/sample_investor.docx already exists")

# ── check other bundled files ─────────────────────────────────────────────────
for fname in ["sample_investor.txt", "sample_investor.json", "sample_investor.csv", "sample_investor.md"]:
    p = DATA_INPUT / fname
    ok(f"data/input/{fname} present") if p.exists() else warn(f"data/input/{fname} MISSING — re-copy from repo")

for fname in ["blank_form.pdf", "sample_investor.pdf"]:
    p = DATA_INPUT / fname
    if p.exists():
        ok(f"data/input/{fname} present")
    else:
        info(f"data/input/{fname} not found — related tests will be skipped")


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — UNIT TESTS
# ══════════════════════════════════════════════════════════════════════════════

header("Step 3 — Unit tests  (no API key, no files needed)")
results = []
r = run_pytest("Unit tests", ["tests/unit/", "-v"])
results.append(r)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — FUNCTIONAL TESTS (mocked LLM)
# ══════════════════════════════════════════════════════════════════════════════

header("Step 4 — Functional tests  (mocked LLM — no API key)")
r = run_pytest(
    "Functional (mocked LLM)",
    ["tests/functional/", "-v",
     "-k", "not real and not mapper_e2e"],
    # Force filler=none so .env's DOC_UPLOAD_PDF_FILLER=mapper doesn't
    # spin up the in-process mapper during mocked extraction tests
    env={"DOC_UPLOAD_PDF_FILLER": "none"},
)
results.append(r)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — REAL LLM TEST
# ══════════════════════════════════════════════════════════════════════════════

header("Step 5 — Real LLM test  (reads API key from .env)")

if api_key and not api_key.startswith("sk-..."):
    r = run_pytest(
        "Real LLM extraction",
        ["tests/functional/", "-v", "-k", "real"],
        env={"DOC_UPLOAD_LLM_API_KEY": api_key},
    )
    results.append(r)
else:
    warn("Skipping — set DOC_UPLOAD_LLM_API_KEY in .env to enable")
    results.append({"label": "Real LLM extraction", "passed": 0, "failed": 0,
                    "skipped": 1, "errors": 0, "ok": True})


# ══════════════════════════════════════════════════════════════════════════════
# STEP 6 — MAPPER END-TO-END
# ══════════════════════════════════════════════════════════════════════════════

header("Step 6 — Mapper end-to-end  (real LLM + real PDF fill)")

blank_pdf_exists = pdf_path and Path(pdf_path).exists()
has_api_key = bool(api_key and not api_key.startswith("sk-..."))

if has_api_key and blank_pdf_exists:
    r = run_pytest(
        "Mapper e2e (txt + json -> PDF fill)",
        ["tests/functional/", "-v", "-k", "mapper_e2e"],
        env={
            "DOC_UPLOAD_LLM_API_KEY": api_key,
            "DOC_UPLOAD_PDF_PATH": pdf_path,
            "DOC_UPLOAD_PDF_FILLER": "mapper",
            "MAPPER_API_URL": "",
        },
    )
    results.append(r)
elif not has_api_key:
    warn("Skipping — set DOC_UPLOAD_LLM_API_KEY in .env to enable")
    results.append({"label": "Mapper e2e", "passed": 0, "failed": 0,
                    "skipped": 1, "errors": 0, "ok": True})
else:
    warn(f"Skipping — blank PDF not found at: {pdf_path!r}")
    warn("Set DOC_UPLOAD_PDF_PATH=<path to blank_form.pdf> in .env")
    results.append({"label": "Mapper e2e", "passed": 0, "failed": 0,
                    "skipped": 1, "errors": 0, "ok": True})


# ══════════════════════════════════════════════════════════════════════════════
# STEP 7 — LOCAL RUNNER SMOKE TEST
# ══════════════════════════════════════════════════════════════════════════════

header("Step 7 — Local runner smoke test  (non-interactive)")

smoke_env = {**os.environ}
if api_key and not api_key.startswith("sk-..."):
    smoke_env["DOC_UPLOAD_LLM_API_KEY"] = api_key

txt_file = DATA_INPUT / "sample_investor.txt"
if txt_file.exists() and has_api_key:
    step("Running: python -m entrypoints.local --document data/input/sample_investor.txt …")
    smoke_env["PYTHONIOENCODING"] = "utf-8"
    smoke = subprocess.run(
        [PYTHON, "-m", "entrypoints.local",
         "--document", str(txt_file),
         "--schema", "configs/form_keys.json"],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        cwd=str(ROOT), env=smoke_env
    )
    if smoke.returncode == 0:
        # Check key outputs
        out = smoke.stdout + smoke.stderr
        if "PIPELINE COMPLETE" in out:
            ok("Pipeline completed successfully")
        if "PDF filled:" in out:
            for line in out.splitlines():
                if "PDF filled:" in line and "✅" in line:
                    ok(line.strip().lstrip("✅ ").strip())
        if "Errors     : 0" in out or "errors\": 0" in out or '"errors": []' in out:
            ok("0 errors")
        results.append({"label": "Local runner smoke test", "passed": 1, "failed": 0,
                        "skipped": 0, "errors": 0, "ok": True})
    else:
        warn("Local runner returned non-zero exit code")
        # Show last 20 lines of output
        for line in (smoke.stdout + smoke.stderr).splitlines()[-20:]:
            print(f"    {line}")
        results.append({"label": "Local runner smoke test", "passed": 0, "failed": 1,
                        "skipped": 0, "errors": 0, "ok": False})
else:
    if not has_api_key:
        warn("Skipping smoke test — no API key in .env")
    else:
        warn("Skipping smoke test — data/input/sample_investor.txt missing")
    results.append({"label": "Local runner smoke test", "passed": 0, "failed": 0,
                    "skipped": 1, "errors": 0, "ok": True})


# ══════════════════════════════════════════════════════════════════════════════
# FINAL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

header("FINAL SUMMARY")

total_passed  = sum(r["passed"]  for r in results)
total_failed  = sum(r["failed"]  for r in results)
total_skipped = sum(r["skipped"] for r in results)
total_errors  = sum(r["errors"]  for r in results)
all_ok = all(r["ok"] for r in results)

print(f"\n  {'Suite':<40} {'Passed':>7} {'Failed':>7} {'Skipped':>8} {'Status':>8}")
print(f"  {'-'*40} {'-'*7} {'-'*7} {'-'*8} {'-'*8}")
for r in results:
    status = "✅ OK" if r["ok"] else "❌ FAIL"
    print(f"  {r['label']:<40} {r['passed']:>7} {r['failed']:>7} {r['skipped']:>8} {status:>8}")
print(f"  {'-'*40} {'-'*7} {'-'*7} {'-'*8} {'-'*8}")
print(f"  {'TOTAL':<40} {total_passed:>7} {total_failed:>7} {total_skipped:>8} {'✅ ALL OK' if all_ok else '❌ FAILURES':>8}")

print()
if total_failed > 0 or total_errors > 0:
    print("  ❌ Some tests failed — check output above for details.")
    sys.exit(1)
else:
    print("  🎉 All tests passed!")
    sys.exit(0)